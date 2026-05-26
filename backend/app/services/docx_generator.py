from datetime import datetime
from typing import Optional
from docx import Document
from docx.shared import Pt

from app.models.profile import Profile
from app.schemas.job_offer import ParsedJobOffer


class DOCXGenerator:
    def __init__(self, template_path: str):
        self.template_path = template_path

    def generate(
        self,
        profile: Profile,
        job_parsed: Optional[ParsedJobOffer],
        output_path: str,
    ):
        doc = Document(self.template_path)

        self._fill_header(doc, profile)

        sections = self._find_sections(doc)
        for heading_para, next_para in reversed(sections):
            heading_text = heading_para.text.strip().lower()

            if "education" in heading_text and "leadership" not in heading_text:
                self._replace_section(doc, heading_para, next_para, profile.education_entries, self._build_education)
            elif heading_text == "experience":
                self._replace_section(doc, heading_para, next_para, profile.experiences, self._build_experience)
            elif "leadership" in heading_text or "activities" in heading_text:
                self._replace_section(doc, heading_para, next_para, profile.projects if profile.projects else None, self._build_project)
            elif "skills" in heading_text:
                self._replace_skills_section(doc, heading_para, next_para, profile, job_parsed)

        for heading_para, _ in sections:
            heading_text = heading_para.text.strip().lower()
            if "leadership" in heading_text or "activities" in heading_text:
                if profile.projects:
                    heading_para.clear()
                    run = heading_para.add_run("Projects")
                    run.bold = True
                else:
                    heading_para._element.getparent().remove(heading_para._element)
            elif "skills" in heading_text:
                heading_para.clear()
                run = heading_para.add_run("Skills")
                run.bold = True

        doc.save(output_path)

    def _find_sections(self, doc: Document) -> list[tuple]:
        heading_map = {
            "education": "Education",
            "experience": "Experience",
            "leadership": "Leadership & Activities",
            "activities": "Leadership & Activities",
            "skills": "Skills & Interests",
        }
        skip_keywords = [
            "beginning with", "begin each", "quantify", "do not use",
            "with your next", "this section", "if this section",
            "technical:", "language:", "laboratory:", "interests:",
            "relevant coursework", "study abroad", "high school", "harvard university",
        ]
        heading_paras = []
        for p in doc.paragraphs:
            text = p.text.strip().lower()
            if not text:
                continue
            matched_keyword = None
            for kw in heading_map:
                if kw in text:
                    matched_keyword = kw
                    break
            if not matched_keyword:
                continue
            if any(sk in text for sk in skip_keywords):
                continue
            heading_paras.append(p)

        sections = []
        for idx, h in enumerate(heading_paras):
            n = heading_paras[idx + 1] if idx + 1 < len(heading_paras) else None
            sections.append((h, n))
        return sections

    def _para_index(self, doc: Document, para) -> int:
        el = para._element
        for i, p in enumerate(doc.paragraphs):
            if p._element is el:
                return i
        return -1

    def _remove_between(self, doc: Document, start_para, end_para):
        start_i = self._para_index(doc, start_para) + 1
        if end_para:
            end_i = self._para_index(doc, end_para) - 1
        else:
            end_i = len(doc.paragraphs) - 1
        if end_i < start_i:
            return
        total = len(doc.paragraphs)
        for i in range(end_i, start_i - 1, -1):
            if i < total:
                doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)

    def _replace_section(self, doc: Document, heading_para, next_para, entries, build_fn):
        has_content = entries is not None
        self._remove_between(doc, heading_para, next_para)

        if not has_content:
            return

        if next_para:
            ref_para = next_para
        else:
            ref_para = doc.add_paragraph()

        lines = build_fn(entries)
        for line_data in lines:
            para = ref_para.insert_paragraph_before()
            self._apply_line(doc, para, line_data)

        if not next_para:
            ref_para._element.getparent().remove(ref_para._element)

    def _apply_line(self, doc: Document, para, line_data):
        if isinstance(line_data, str):
            run = para.add_run(line_data)
            run.font.name = "Calibri"
        elif isinstance(line_data, dict):
            for segment in line_data.get("runs", []):
                run = para.add_run(segment.get("text", ""))
                run.bold = segment.get("bold", False)
                run.font.name = "Calibri"
            if line_data.get("bullet", False):
                try:
                    para.style = doc.styles["List Paragraph"]
                except Exception:
                    pass

    def _replace_skills_section(self, doc: Document, heading_para, next_para, profile: Profile, job_parsed: Optional[ParsedJobOffer]):
        self._remove_between(doc, heading_para, next_para)

        if not profile.skills and not profile.languages:
            return

        if next_para:
            ref_para = next_para
        else:
            ref_para = doc.add_paragraph()

        hard_skills = [s for s in profile.skills if s.category in ("hard", "technical")]
        soft_skills = [s for s in profile.skills if s.category == "soft"]

        items = []
        if hard_skills:
            items.append(("Technical: ", ", ".join(s.name for s in hard_skills)))
        if soft_skills:
            items.append(("Professional: ", ", ".join(s.name for s in soft_skills)))
        if profile.languages:
            lang_texts = [f"{l.language_name} ({l.proficiency})" for l in profile.languages]
            items.append(("Languages: ", ", ".join(lang_texts)))

        for prefix, text in items:
            para = ref_para.insert_paragraph_before()
            run = para.add_run(prefix)
            run.bold = True
            run.font.name = "Calibri"
            run = para.add_run(text)
            run.font.name = "Calibri"

        if not next_para:
            ref_para._element.getparent().remove(ref_para._element)

    def _fill_header(self, doc: Document, profile: Profile):
        if len(doc.paragraphs) > 1:
            name_para = doc.paragraphs[1]
            name_para.clear()
            run = name_para.add_run(profile.full_name or "Your Name")
            run.bold = True
            run.font.size = Pt(16)
            run.font.name = "Calibri"

        contact_parts = []
        if profile.location:
            contact_parts.append(profile.location)
        if profile.email:
            contact_parts.append(profile.email)
        if profile.phone:
            contact_parts.append(profile.phone)
        contact_text = " | ".join(contact_parts)
        if profile.linkedin_url:
            contact_text += f" | {profile.linkedin_url}"
        if profile.portfolio_url:
            contact_text += f" | {profile.portfolio_url}"

        for p in doc.paragraphs:
            pt = p.text.strip()
            if pt and ("street" in pt.lower() or "youremail" in pt.lower() or "@" in pt.lower() and "center" not in pt.lower()):
                p.clear()
                run = p.add_run(contact_text)
                run.font.size = Pt(9)
                run.font.name = "Calibri"
                break

    def _build_education(self, entries) -> list:
        lines = []
        for edu in entries:
            lines.append({
                "runs": [
                    {"text": edu.institution, "bold": True},
                ]
            })

            details = []
            if edu.degree:
                details.append(edu.degree)
            if edu.field_of_study:
                details.append(edu.field_of_study)
            if edu.gpa:
                details.append(f"GPA: {edu.gpa}")

            date_str = ""
            if edu.start_date:
                date_str = self._fmt_date(edu.start_date)
            if edu.end_date:
                date_str += f" – {self._fmt_date(edu.end_date)}"

            detail_text = ", ".join(details)
            lines.append({
                "runs": [
                    {"text": detail_text, "bold": False},
                    {"text": f"\t{date_str}"},
                ]
            })

            if edu.honors:
                lines.append(f"Honors: {', '.join(edu.honors)}")

        return lines

    def _build_experience(self, entries) -> list:
        sorted_exps = sorted(
            entries,
            key=lambda e: e.start_date or datetime.min.date(),
            reverse=True,
        )
        lines = []
        for exp in sorted_exps:
            lines.append({
                "runs": [
                    {"text": exp.company, "bold": True},
                    {"text": f"\t{exp.location or ''}"},
                ]
            })

            date_str = ""
            if exp.start_date:
                date_str = self._fmt_date(exp.start_date)
            date_str += " – "
            if exp.is_current:
                date_str += "Present"
            elif exp.end_date:
                date_str += self._fmt_date(exp.end_date)

            lines.append({
                "runs": [
                    {"text": exp.title, "bold": True},
                    {"text": f"\t{date_str}"},
                ]
            })

            for bullet in (exp.bullets or []):
                lines.append({"runs": [{"text": bullet, "bold": False}], "bullet": True})

        return lines

    def _build_project(self, entries) -> list:
        lines = []
        for project in entries:
            lines.append({
                "runs": [
                    {"text": project.name, "bold": True},
                ]
            })

            if project.technologies:
                lines.append(f"Technologies: {', '.join(project.technologies)}")

            for bullet in (project.bullets or []):
                lines.append({"runs": [{"text": bullet, "bold": False}], "bullet": True})

            if not project.bullets and project.description:
                lines.append(project.description)

        return lines

    def _fmt_date(self, d) -> str:
        if hasattr(d, "strftime"):
            return d.strftime("%b %Y")
        return str(d)
